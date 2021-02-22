"""
This is the file to be executed. It contains most of the higher-level logic.
"""

from docx import Document
from docx.opc import exceptions as docexc
import os
from formatter import Formatter
from spectrum import Spectrum
from exceptions import InputError
from block import Block
import sys
from msvcrt import getch

# Enclose this in comments to enable debugging mode, starting here...


def show_exception_and_exit(exc_type, exc_value, tb):
    import traceback
    traceback.print_exception(exc_type, exc_value, tb)
    print("Press any key to exit.")
    getch()
    sys.exit(-1)


sys.excepthook = show_exception_and_exit
sys.tracebacklimit = 0

# ...right up to here. I'm not sure this works the way it should, though,
# it still exits the interpreter before anything can be read.


PATH = os.path.dirname(os.path.realpath(__file__)) + '/'

def read_document():
    """Access the document with NMR spectra. This only accepts .docx files"""

    global filename
    # This ensures the value for the `filename` variable can be accessed from
    # outside the function, even though the actual value is given within the
    # function.

    filename = input('Enter the input file name: ')
    try:
        document = Document(PATH + filename + '.docx')
    except docexc.PackageNotFoundError:
        print('Error: no document with such name has been found!')
        raise
    else:
        return document


def read_format_style(paragraphs, anchor):
    """
    Read the format style template for the spectra to be read/written.

    :param paragraphs: the spectra .docx document parsed into separate
    paragraphs.
    :param anchor: the keyword found at the beginning of a paragraph
    containing the necessary format template.

    :return: a `Formatter` object describing the format style for a spectrum.
    """

    for paragraph in paragraphs:
        if anchor in paragraph.text:
            raw_format = Block(paragraph, runs=True)[len(anchor):]
            # Clip the keyword and the space after it to be processed into a
            # `Formatter` object.
            return Formatter(raw_format)
    raise InputError('format template not located within the file')


def read_spectra(paragraphs, formatter):
    """
    Read the spectra data from a list of .docx paragraphs and return a
    list of `Spectrum` objects.

    :param paragraphs: a list of `Paragraph` objects.
    :param formatter: a `Formatter` object used to parse the raw data.

    :return: a list of `Spectrum` objects.
    """

    r_spectra = []
    for i, paragraph in enumerate(paragraphs):
        if paragraph.text.split(' ')[0] == 'Spectrum:':
        # This keyword signals that the next paragraphs contains the actual
        # spectrum data.

            r_spectra.append((paragraph.text.rstrip().replace('Spectrum: ',
                ''),
                              paragraphs[i+1]))
            # Store the cypher given after the keyword and the spectrum
            # itself as a tuple of raw data, itself appended to a list of
            # raw data.

    print('Number of located spectra: {}.\n'.format(len(r_spectra)))

    spectra = []    # The list to store processed spectra.
    for r_spectrum in r_spectra:
        spectrum = Spectrum(r_spectrum, formatter) # Build a `Spectrum`
                                                   # object for each tuple
                                                   # in the list of raw data.
        print(spectrum)
        spectra.append(spectrum)    # Save each `Spectrum` object into a list.
        for signal in spectrum.signals:
            print(signal)
        print('\n')
    return spectra


def read_task(paragraphs):
    """
    Locate which tasks are to be performed with the spectra.
    :param paragraphs: paragraphs that contain a list of tasks to be performed
    :return: a list of `str` tasks to be performed
    """

    for paragraph in paragraphs:
        if paragraph.text.split(' ')[0] == 'Task:':
            return paragraph.text.split(' ')[1:]
        # Find the keyword preceding listed tasks and return a list of `str`
        # tasks.


def write_spectra(spectra, formatter, document):
    """
    Write every `Spectrum` in a given list to a .docx document with the
    correct styling.

    :param spectra: a list of `Spectrum` data to be written
    :param formatter: `Formatter` object used to write the spectra.
    :param document: the .docx file for writing the spectra down.
    :return: `None`
    """
    for spectrum in spectra:
        spectrum.write(document, formatter)


def ready_reassignment_spectrum(cypher, paragraphs):
    """
    This function creates an empty `Spectrum` that only contains signal
    assignments. The function reads the reassignment values, builds a `str`
    object from it and feeds it to the `Spectrum` class to initiate a
    spectrum. ONLY ONE REASSIGNMENT `Spectrum` WILL BE CREATED!

    :param cypher: `str`, defines which spectrum is to be reassigned. If '*'
    is given, use these reassigns for every spectrum.

    :param paragraphs: a list of paragraphs each containing a reassignment
    value.

    :return: `tuple` containing a list of old assigns for reference and a
    `Spectrum` with new assign values.
    """

    old_assigns = []
    # This is where the assign values to be written over are registered

    new_assigns = Block('')
    # A `Block` object to contain each signals new assignment.

    document = Document()
    document.add_paragraph()
    # The `__init__` function in the `Spectrum` class requires data being
    # given as a tuple of `str` cypher and a paragraphs containing the
    # actual spectrum data, so we create that.

    for paragraph in paragraphs:
        if ' = ' not in paragraph.text:
            break
            # The moment there's no ' = ' in a paragraph, the search for
            # reassignment values is over, time to build the reassignment
            # `Spectrum`.
        old_assign = paragraph.text[:paragraph.text.index(' = ')]
        old_assigns.append(old_assign)
        # Everything preceding ' = ' is an old assignment.

        new_assign = Block(paragraph,
            runs=True)[paragraph.text.index(' = ') + 3:]
        new_assigns.append(new_assign)
        # And everything after that is the new assignment. Using the `Block`
        # class to read the new assignment allows to save its styling.
        # Appending this block to `new_assigns` (also a `Block`) results in
        # one big `Block` where the styling of each `Char` is specified.

        new_assigns.append(Block('$'))
        # The '$' special symbol is used as a delimiter between new
        # assignments.

    del new_assigns[-1]
    # The final '$' symbol is unnecessary and will cause problem. We delete it.

    formatter = Formatter(Block(' /%a/$/'))
    # Build a `Formatter` object for parsing the newly built `new_assigns`
    # Block. The first space is necessary, because there always must be a
    # space before the first slash.

    return old_assigns, Spectrum((cypher, new_assigns), formatter,
        blocked_signals_only=True)
    # Using 'blocked_signals_only=True' ensures the initation of an object
    # won't try to make a new Block nor will it attempt to parse head and
    # will consider everything a list of signals instead. Due to the fact
    # that `Spectrum` stores all of its signals in an ordered list (
    # corresponding to the order in which they were parsed), the order
    # between `old_assigns` and signals in the reassignment `Spectrum`  stays
    # the same.



def reassign_spectra(spectra, reassigns):
    """
    This function collectively attempts to perform specific reassignment
    of all the spectra fed to it.

    :param spectra: a `tuple` containing old assignments and the `Spectrum`
    data to be reassigned.
    :param reassigns: a list of `Spectrum` data containing signals to reassign.
    :return: a `Spectrum` object, containing non-specific reassigns for every
    spectrum in the data. If no such global reassignment `Spectrum` is
    provided, return `None`
    """

    ultra_reassign = None
    for new_assign in reassigns:
        if new_assign[1].cypher == '*':
            ultra_reassign = new_assign
            continue
            # The global, non-specific reassign spectrum is the one whose
            # cypher is '*'

        for spectrum in spectra:
            if spectrum.cypher == new_assign[1].cypher:
                reassign_spectrum(spectrum, new_assign)
            # If cyphers of the compared `Spectrum` and reassignment
            # `Spectrum` are the same, perform reassignment.

    return ultra_reassign



def reassign_spectrum(spectrum, new_assign):
    """
    The actual function performing reassignment.

    :param spectrum: `Spectrum` to be reassigned.
    :param new_assign: a `tuple` of corresponding old assignments and
    reassignment `Spectrum`.
    :return: None
    """

    old_assignments = [x.assignment() for x in spectrum.signals]
    # Make a list of signals in the original spectrum.

    for j, signal in enumerate(new_assign[0]):
        if signal in old_assignments:
            i = old_assignments.index(signal)
            spectrum.signals[i].assignment =\
                new_assign[1].signals[j].assignment
            # If there's a match between a signal to be reassigned and a
            # signal found in the spectrum, then locate the index of the
            # signal to be reassigned within the processed `Spectrum`
            # (`spectrum.signals[i].assignment`), and substitute it for the
            # new value (`new_assign[-1].signals[j].assignment). Since all of
            # this is done with `Block` objects, the styling is retained.



filename = ''
# This sets up a global variable later to be used in `read_document` function.

document = read_document()
paragraphs = document.paragraphs
formatter = read_format_style(paragraphs, 'Input format: ')
spectra = read_spectra(paragraphs, formatter)
# Read the document, extract the paragraphs, find the input format there
# using the keyword and use the built `Formatter` object to parse spectral
# data.

task = read_task(paragraphs)
# Read the tasks.

if 'reassign' in task:
    reassigns = []
    for i, paragraph in enumerate(paragraphs):
        if paragraph.text.split(' ')[0] == 'Assignments:':
            cypher = paragraph.text.replace('Assignments: ', '')
            reassigns.append(ready_reassignment_spectrum(cypher,
                paragraphs[i + 1:]))
    ultra_reassign = reassign_spectra(spectra, reassigns)
    if ultra_reassign:
        for spectrum in spectra:
            reassign_spectrum(spectrum, ultra_reassign)
# If the task is to reassign the old assignments, then first it's necessary
# to find the paragraph that starts with the keyword (and cut it out).
# Everything after the keyword is the cypher of the spectrum to be
# reassigned (or '*' if its the non-specific reassignment). Then
# feed every paragraph after that to `ready_reassignment_spectrum` that will
# return a tuple of old assignments and a `Spectrum` with new assignments.
# Append it to the `reassigns` list. Then perform specific reassignment for
# all spectra with `reassign_spectra`. If there's a non-specific
# reassignment to be done, `ultra_reassign` will contain a tuple rather than
# None, feed it to `reassign_spectrum` for every spectrum to perform
# non-specific reassignment. Note that specific reassignment has to take
# place first, since non-specific reassignment might erroneously reassign a
# signal before the specific reassignment gets a chance to do so.
# Thought: specific reassignment might unintentionally feed the right
# conditions for non-specific reassignment, but since thus far I only needed
# to reassign minimalistic notation to actual assignment, this has not been
# a problem.

if 'convert' in task:
    formatter = read_format_style(paragraphs, 'Output format: ')
# If one needs to use a new format for writing the spectra down, ready the
# new `Formatter` object.


while True:
    new_file = input('Enter the name of the file to be written: ')
    if new_file == filename:
        print('You are trying to re-write the input file! Please, choose '
              'another name for your output file!')
        continue
    break
new_document = Document()
write_spectra(spectra, formatter, new_document)
new_document.save(PATH + new_file + '.docx')
# Ask where to write the new file. Saving atop the same file as the input
# file is forbidden. Perform `write_spectra` using the wanted `Formatter`
# and save the .docx file.

os.system('pause')
# This is intended to let the user see the Error Message should any arise,
# but this doesn't actually work like this :'(.
